from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from saida.ingestion.parsers import detect_kind, parse_docx, parse_pdf, parse_text, parse_xlsx


def test_detect_kind_marks_xlsx_as_tabular():
    assert detect_kind("sample.xlsx") == "tabular"


def test_parse_docx_extracts_paragraph_text(tmp_path: Path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Quarterly revenue analysis")
    doc.add_paragraph("Margins improved in Q4")
    doc.save(path)

    text = parse_docx(str(path))
    assert "Quarterly revenue analysis" in text
    assert "Margins improved in Q4" in text


def test_parse_xlsx_extracts_sheet_rows(tmp_path: Path):
    path = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Revenue"
    ws.append(["quarter", "revenue"])
    ws.append(["Q1", 100])
    ws.append(["Q2", 150])
    wb.save(path)

    text = parse_xlsx(str(path))
    assert "[Sheet: Revenue]" in text
    assert "quarter\trevenue" in text
    assert "Q1\t100" in text


def test_parse_pdf_handles_pdf_file(tmp_path: Path):
    path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as f:
        writer.write(f)

    text = parse_pdf(str(path))
    assert isinstance(text, str)


def test_parse_text_dispatches_by_extension(tmp_path: Path):
    docx_path = tmp_path / "dispatch.docx"
    doc = Document()
    doc.add_paragraph("Dispatch check")
    doc.save(docx_path)

    assert "Dispatch check" in parse_text(str(docx_path))
